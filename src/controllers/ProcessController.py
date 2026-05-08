"""
ProcessController — loads files and splits them into text chunks.

Supported formats: .txt, .pdf, .docx, .html
All formats go through Arabic text normalization before chunking.
"""

import os
import re
import unicodedata
import chardet
from langchain_community.document_loaders import TextLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document

from controllers.BaseController import BaseController
from controllers.FileController import FileController


class ProcessController(BaseController):
    """Loads files into Documents and splits them into chunks."""

    def __init__(self, project_id: str):
        super().__init__()
        file_controller = FileController()
        self.project_path = file_controller.get_file_path(project_id)

    @staticmethod
    def get_file_extension(file_name: str) -> str:
        """Return the lowercase file extension (e.g. '.pdf')."""
        return os.path.splitext(file_name)[-1].lower()

    def get_file_loader(self, file_name: str):
        """Pick the right loader based on file extension."""
        file_path = os.path.join(self.project_path, file_name)
        ext = self.get_file_extension(file_name)

        if ext == ".txt":
            detected = self._detect_encoding(file_path)
            return TextLoader(file_path, encoding=detected)

        if ext == ".html":
            return self._load_html(file_path)

        if ext == ".pdf":
            return self._load_pdf(file_path)

        if ext == ".docx":
            return self._load_docx(file_path)

        return None

    def get_file_content(self, file_name: str) -> list[Document] | None:
        """Load a file and return its content as Documents."""
        loader = self.get_file_loader(file_name)

        if loader is None:
            return None

        # Custom loaders return a list directly, TextLoader needs .load()
        if isinstance(loader, list):
            return loader

        try:
            return loader.load()
        except FileNotFoundError:
            return None

    def process_files(
        self,
        file_content: list[Document],
        file_name: str,
        chunk_size: int = 100,
        overlap: int = 20,
    ) -> list[Document]:
        """Normalize text, then split into overlapping chunks."""
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=overlap,
            length_function=len,
        )

        for doc in file_content:
            doc.page_content = self._normalize_arabic_text(doc.page_content)
            doc.metadata = doc.metadata or {}
            doc.metadata["source_file"] = file_name

        chunks = text_splitter.create_documents(
            texts=[doc.page_content for doc in file_content],
            metadatas=[doc.metadata for doc in file_content],
        )

        return chunks

    # --- Loaders ---

    @staticmethod
    def _load_pdf(file_path: str) -> list[Document]:
        """Load PDF using PyMuPDF. sort=True fixes RTL Arabic text order."""
        import pymupdf

        pages = []
        with pymupdf.open(file_path) as pdf_doc:
            for page_num, page in enumerate(pdf_doc):
                text = page.get_text("text", sort=True)
                if text.strip():
                    pages.append(
                        Document(
                            page_content=text,
                            metadata={
                                "source": file_path,
                                "page": page_num + 1,
                            },
                        )
                    )
        return pages

    @staticmethod
    def _load_docx(file_path: str) -> list[Document]:
        """Load DOCX — extracts paragraphs, tables, headers, and footers."""
        from docx import Document as DocxDocument

        docx_doc = DocxDocument(file_path)
        parts = []

        # Paragraphs
        for para in docx_doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)

        # Table cells
        for table in docx_doc.tables:
            for row in table.rows:
                row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_texts:
                    parts.append(" | ".join(row_texts))

        # Headers and footers
        for section in docx_doc.sections:
            for header_para in section.header.paragraphs:
                if header_para.text.strip():
                    parts.append(header_para.text)
            for footer_para in section.footer.paragraphs:
                if footer_para.text.strip():
                    parts.append(footer_para.text)

        return [
            Document(
                page_content="\n".join(parts),
                metadata={"source": file_path},
            )
        ]

    def _load_html(self, file_path: str) -> list[Document]:
        """Load HTML — only extracts text from content tags (p, h1-h6, li, etc.)."""
        from bs4 import BeautifulSoup

        detected = self._detect_encoding(file_path)
        with open(file_path, "r", encoding=detected, errors="replace") as f:
            soup = BeautifulSoup(f, "html.parser")

        content_tags = ["p", "h1", "h2", "h3", "h4", "h5", "h6",
                        "li", "td", "th", "blockquote"]

        paragraphs = []
        for tag in soup.find_all(content_tags):
            text = tag.get_text(separator=" ", strip=True)
            if text and len(text) > 1:
                paragraphs.append(text)

        title = str(soup.title.string) if soup.title else ""

        return [
            Document(
                page_content="\n".join(paragraphs),
                metadata={"source": file_path, "title": title},
            )
        ]

    # --- Helpers ---

    @staticmethod
    def _detect_encoding(file_path: str) -> str:
        """Auto-detect file encoding. Falls back to utf-8 if unsure."""
        with open(file_path, "rb") as f:
            raw = f.read()
        result = chardet.detect(raw)
        encoding = result.get("encoding") or "utf-8"
        confidence = result.get("confidence") or 0

        if confidence < 0.5:
            encoding = "utf-8"

        return encoding

    @staticmethod
    def _normalize_arabic_text(text: str) -> str:
        """
        Clean up Arabic text:
        1. NFKC normalize — converts special ligature forms (e.g. ﻻ)
           back to standard Arabic letters (e.g. لا). PDFs often store
           Arabic using these visual forms instead of real characters.
        2. Strip diacritics (tashkeel) and tatweel
        3. Normalize alef variants (أ إ آ -> ا)
        4. Collapse whitespace
        """
        if not text:
            return text

        # Convert visual/ligature forms to standard Arabic characters
        text = unicodedata.normalize("NFKC", text)

        # Remove diacritics (harakat) and tatweel (stretching)
        text = re.sub(r"[\u064B-\u065F\u0670\u0640]", "", text)

        # Normalize alef variants to plain alef
        text = re.sub(r"[\u0622\u0623\u0625]", "\u0627", text)

        # Collapse whitespace
        text = re.sub(r"\s+", " ", text)

        return text.strip()
